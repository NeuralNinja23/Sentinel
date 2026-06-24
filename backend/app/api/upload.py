import io
import logging
from fastapi import APIRouter, UploadFile, File, HTTPException, status
from pydantic import BaseModel
from typing import List

import pypdf
import docx

logger = logging.getLogger("upload")
router = APIRouter(prefix="/api/upload", tags=["upload"])

# Global in-memory list to store parsed documents.
# Structure: {"filename": str, "content": str, "size": int}
uploaded_documents = []

class DocumentInfo(BaseModel):
    filename: str
    size: int  # file size in bytes

def extract_pdf_text(file_bytes: bytes) -> str:
    try:
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text_content = []
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Could not parse PDF file: {str(e)}"
        )

def extract_docx_text(file_bytes: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text_content = []
        # Extract paragraph text
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_content.append(paragraph.text)
        # Extract table text if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_content.append(cell.text)
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Could not parse DOCX file: {str(e)}"
        )

def extract_text_file(file_bytes: bytes) -> str:
    # Try different encodings to decode text cleanly
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
    for encoding in encodings:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail="Could not decode text file with standard encodings."
    )

@router.post("", response_model=List[DocumentInfo])
async def upload_document(file: UploadFile = File(...)):
    filename = file.filename
    if not filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Filename is missing."
        )
    
    lower_filename = filename.lower()
    # Check allowed text documents, strictly blocking images and other formats
    if not (lower_filename.endswith(".pdf") or 
            lower_filename.endswith(".docx") or 
            lower_filename.endswith(".txt") or 
            lower_filename.endswith(".md")):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported file format. Only PDF, DOCX, TXT, and MD files are allowed."
        )
        
    try:
        file_bytes = await file.read()
        file_size = len(file_bytes)
        
        # Parse based on extension
        if lower_filename.endswith(".pdf"):
            content = extract_pdf_text(file_bytes)
        elif lower_filename.endswith(".docx"):
            content = extract_docx_text(file_bytes)
        else:  # txt or md
            content = extract_text_file(file_bytes)
            
        # Check if file with same name is already uploaded, remove it to overwrite
        uploaded_documents[:] = [doc for doc in uploaded_documents if doc["filename"] != filename]
        
        # Append parsed details
        uploaded_documents.append({
            "filename": filename,
            "content": content,
            "size": file_size
        })
        
        logger.info(f"Successfully uploaded and parsed document: {filename} ({file_size} bytes)")
        
        return get_current_documents_info()
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error during file upload: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"An unexpected error occurred while processing the file: {str(e)}"
        )

@router.post("/clear", response_model=List[DocumentInfo])
async def clear_documents():
    uploaded_documents.clear()
    logger.info("Cleared all uploaded documents context.")
    return []

@router.delete("/{filename}", response_model=List[DocumentInfo])
async def delete_document(filename: str):
    uploaded_documents[:] = [doc for doc in uploaded_documents if doc["filename"] != filename]
    logger.info(f"Deleted document context: {filename}")
    return get_current_documents_info()

@router.get("/list", response_model=List[DocumentInfo])
async def list_documents():
    return get_current_documents_info()

def get_current_documents_info() -> List[DocumentInfo]:
    return [
        DocumentInfo(filename=doc["filename"], size=doc["size"])
        for doc in uploaded_documents
    ]

def get_all_documents_text_context() -> str:
    """Helper to retrieve consolidated text from all uploaded documents for context injection."""
    if not uploaded_documents:
        return ""
    
    context_blocks = []
    for doc in uploaded_documents:
        context_blocks.append(
            f"=== DOCUMENT: {doc['filename']} ===\n"
            f"{doc['content']}\n"
            f"=== END OF DOCUMENT ==="
        )
    return "\n\n".join(context_blocks)
